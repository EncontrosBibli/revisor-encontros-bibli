import streamlit as st
from docx import Document
import google.generativeai as genai
import time
from io import BytesIO

# --- 1. CONFIGURAÇÃO DA PÁGINA ---
st.set_page_config(
    page_title="Editoria Encontros Bibli", 
    layout="wide", 
    page_icon="logo_revista.png" 
)

# --- 2. ESTILIZAÇÃO CSS (O visual que você aprovou) ---
st.markdown("""
    <style>
    .stApp { background-color: #f0f2f5; }
    [data-testid="stSidebar"] { background-color: #70298d; }
    [data-testid="stSidebar"] * { color: white !important; }

    .main-header {
        background-color: #70298d;
        padding: 25px;
        border-radius: 0px 0px 10px 10px;
        color: white !important;
        margin-bottom: 25px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    }
    
    .main-header h1, .main-header p { color: white !important; margin: 0; }

    .logo-eb {
        background-color: white; color: #70298d; padding: 5px 12px; 
        border-radius: 6px; margin-right: 15px; 
        border-left: 6px solid #ff8c00; box-shadow: 3px 3px 0px #ff8c00;
        font-family: 'Arial Black', sans-serif;
    }

    h1, h2, h3 { color: #31333f !important; font-family: 'Segoe UI', sans-serif; }

    .stButton>button {
        width: 100%; border-radius: 8px;
        background-color: #70298d; color: white !important;
        font-weight: bold; border: none;
    }

    .stButton>button:hover {
        background-color: #5a2172; border: 1px solid #ff8c00; color: #ff8c00 !important;
    }

    .stTabs { background-color: #ffffff; padding: 20px; border-radius: 12px; }
    </style>
    """, unsafe_allow_html=True)

# Cabeçalho com o EB estilizado
st.markdown("""
    <div class="main-header">
        <h1><span class="logo-eb">EB</span> Editor de Encontros Bibli</h1>
        <p>Sistema de Revisão Técnica - Sincronizado com Tutorial 10/04/2025</p>
    </div>
    """, unsafe_allow_html=True)

# --- 3. LOGO SIDEBAR ---
try:
    st.sidebar.image("logo_revista.png", use_container_width=True)
except:
    st.sidebar.write("## Encontros Bibli")

st.sidebar.markdown("---")

# --- 4. FUNÇÕES DE APOIO (Lógica que deu certo) ---
def gerar_docx(conteudo, titulo):
    doc = Document()
    doc.add_heading(titulo, 0)
    for linha in conteudo.split('\n'):
        if linha.strip(): doc.add_paragraph(linha)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def realizar_analise(prompt, api_key):
    try:
        genai.configure(api_key=api_key)
        modelos_validos = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        nome_modelo = next((m for m in modelos_validos if 'gemini-1.5-flash' in m), modelos_validos[0])
        model = genai.GenerativeModel(nome_modelo)
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        if "429" in str(e): return "⚠️ ERRO_COTA: Limite atingido. Aguarde 60s."
        return f"❌ Erro na API: {str(e)}"

# --- 5. INTERFACE ---
with st.sidebar:
    st.header("Configuração")
    api_key_input = st.text_input("🔑 API Key:", type="password")
    api_key = api_key_input if api_key_input else st.secrets.get("GEMINI_API_KEY", "")
    if st.button("🧹 Limpar Sessão"):
        st.session_state.clear()
        st.rerun()

if not api_key:
    st.info("👈 Por favor, insira sua API Key na lateral para ativar o sistema.")
    st.stop()

artigo_file = st.file_uploader("📂 Subir Artigo para Revisão (DOCX)", type="docx")

if artigo_file:
    with st.spinner("⏳ Lendo artigo e mapeando diretrizes da UFSC..."):
        try:
            doc = Document(artigo_file)
            texto_artigo = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
            st.success(f"✅ Artigo '{artigo_file.name}' carregado!")
        except Exception as e:
            st.error(f"Erro: {e}"); st.stop()

    tab1, tab2, tab3 = st.tabs(["📐 Maquetação & Forma", "✍️ Gramática & Citações", "📚 Referências (ABNT)"])

# --- ABA 1: MAQUETAÇÃO (Atualizada com Tipologia) ---
    with tab1:
        st.subheader("Análise de Maquetação, Formatação e Ilustrações")
        if st.button("Executar Análise de Forma"):
            with st.spinner("Conferindo normas, ilustrações e tipologia..."):
                prompt = (
                    "Aja como Editor de Layout da Revista Encontros Bibli. Analise o artigo conforme o Tutorial de 10/04/2025 e o Template oficial:\n\n"
                    "1. TIPOLOGIA (Cabeçalho da 1ª página):\n"
                    "   - Verificar: Artigo original, Artigo de dados, Ensaio ou Estudo de casos.\n"
                    "   - Traduzir a tipologia se o artigo for em EN ou ES.\n\n"
                    "2. CORPO DO TEXTO (Formatação):\n"
                    "   - Alinhamento: Deve ser JUSTIFICADO.\n"
                    "   - Recuo: Início de parágrafo deve ter recuo de 1,25 cm.\n"
                    "   - Espaçamento: Entre linhas deve ser 1,5.\n"
                    "   - Verifique se há espaços duplos desnecessários entre parágrafos.\n\n"
                    "3. ILUSTRAÇÕES (Tabelas, Quadros e Figuras):\n"
                    "   - Referenciação: Devem ser citadas no corpo do texto antes de aparecerem (ex: 'conforme a Tabela 1...').\n"
                    "   - Identificação: Título acima da ilustração (ex: Tabela 1 – Título).\n"
                    "   - Fontes: Devem ter a indicação da fonte logo abaixo da ilustração (ex: Fonte: Dados da pesquisa (2025)).\n\n"
                    "4. TÍTULOS E METADADOS:\n"
                    "   - Título PT: Arial Black, 16, MAIÚSCULO, Negrito.\n"
                    "   - Título EN: Arial, 10, minúsculo, Negrito.\n"
                    "   - Palavras-chave: 3 a 5 termos separados por PONTO (.).\n"
                    "   - Remover menções a 'uso exclusivo'.\n"
                    f"\nTexto para análise:\n{texto_artigo[:15000]}"
                )
                res = realizar_analise(prompt, api_key)
                st.markdown(res)
                st.download_button("📥 Baixar Relatório", gerar_docx(res, "Revisao_Maquetacao"), "maquetacao.docx")
    with tab2:
        st.subheader("Revisão Linguística, Citações e Consistência")
        if st.button("Executar Revisão Linguística"):
            with st.spinner("Analisando gramática e cruzando citações com referências..."):
                # Capturamos o meio do texto (citações) e o fim (referências) para o cruzamento
                corpo_texto = texto_artigo[1000:15000]
                lista_referencias = texto_artigo[-8000:]
                
                prompt = (
                    "Aja como revisor acadêmico sênior da Encontros Bibli. Sua tarefa é dupla:\n\n"
                    "1. REVISÃO LINGUÍSTICA E NBR 10520:\n"
                    "   - Analise gramática, ortografia e estilo científico.\n"
                    "   - Verifique citações longas (+3 linhas): Devem ter recuo 4cm, Arial 10, sem aspas.\n"
                    "   - Garante que 'et al.' esteja em itálico.\n\n"
                    "2. CRUZAMENTO DE CITAÇÕES (CRÍTICO):\n"
                    "   - Verifique se todos os autores citados no corpo do texto aparecem na lista de referências final.\n"
                    "   - Identifique autores que estão nas referências mas não foram citados no texto.\n"
                    "   - Liste explicitamente as ausências (Ex: 'Silva (2022) citado no texto, mas ausente nas referências').\n"
                    f"\n--- CORPO DO TEXTO ---\n{corpo_texto}\n"
                    f"\n--- LISTA DE REFERÊNCIAS ---\n{lista_referencias}"
                )
                res = realizar_analise(prompt, api_key)
                st.markdown(res)
                st.download_button("📥 Baixar Relatório", gerar_docx(res, "Revisao_Gramatical_Consistencia"), "revisao_linguistica.docx")

    with tab3:
        st.subheader("Validação de Referências (NBR 6023)")
        if st.button("Validar Referências"):
            with st.spinner("Conferindo ABNT..."):
                prompt = (
                    "Aja como Editor da Encontros Bibli. Valide conforme NBR 6023:\n"
                    "1. Título da obra em NEGRITO.\n"
                    "2. DOI em formato de URL completa (https://doi.org/...).\n"
                    "3. Padrão SOBRENOME, Nome.\n"
                    f"\nReferências:\n{texto_artigo[-8000:]}"
                )
                res = realizar_analise(prompt, api_key)
                st.markdown(res)
                st.download_button("📥 Baixar Relatório", gerar_docx(res, "Referencias"), "referencias.docx")













